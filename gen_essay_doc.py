#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成作文 Word（润色版）：当《简爱》走上银幕"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('当《简爱》走上银幕')
r.font.size = Pt(16)
r.font.bold = True
r.font.name = 'Times New Roman'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title.paragraph_format.space_after = Pt(18)

paragraphs = [
    '翻开书页，简爱从纸上走来；走进影院，简爱从光里走来。读完原著再看电影，两个简爱在我心里重逢，也让我忍不住追问：电影对书籍的改编，究竟是在成全一本书，还是在改写一本书？',
    '电影自有书籍比不上的魅力。当简爱在薄雾中第一次遇见坠马受伤的罗切斯特，当桑菲尔德深夜燃起大火，光影与音乐一下子把书中的文字变成了眼前的世界。米娅·华希科沃斯卡演活了那个倔强而自尊的简爱，让我觉得，书里那个“贫穷、低微、不美”的女孩，真的走到了我的面前。这种直观的感动，是文字难以做到的。',
    '可是，看电影时我总觉得少了点什么。原著里有太多“看不见”的东西：简爱在洛伍德八年的成长，格蕾丝·普尔那条神秘的线索，还有罗切斯特最后渐渐恢复视力——那个温暖得让人落泪的结局，在电影里被轻轻删去了。书可以慢慢地讲，电影却只有两个小时，于是改编往往成了一场“减法”，删掉的常常是藏在字里行间的温度。',
    '那么，我喜欢这样的改编吗？我的答案是：喜欢，但更喜欢原著。电影像一位热情的向导，用两个小时带我走近《简爱》；原著则像一位耐心的朋友，用三十八章陪我慢慢长大。先看电影，再读原著，我反而更能体会简爱那句话的分量——“我们的灵魂是平等的。”',
    '书里的简爱住在文字里，银幕上的简爱活在光影中，而真正的简爱，其实住在每一个读者的心里。电影用两个小时讲述她，原著用三十八章陪伴她，它们从来不是对手，而是同一条路上两盏不同的灯。我会先走进影院，再翻开书页——在两个世界里，与同一个不屈的灵魂，一遍又一遍地相遇。',
]

for text in paragraphs:
    para = doc.add_paragraph(text)
    para.paragraph_format.first_line_indent = Pt(24)
    para.paragraph_format.space_after = Pt(0)

out = '/home/zyt/.openclaw/workspace/作文_当简爱走上银幕.docx'
doc.save(out)
print('saved:', out)

total = sum(len(t) for t in paragraphs)
print('含标点总字数:', total)
